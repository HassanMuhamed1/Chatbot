import os
from typing import TypedDict, List, Union
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.prompts import PromptTemplate
from langchain.schema import HumanMessage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from PIL import Image, ImageFilter, ImageEnhance
import zipfile
from io import BytesIO
import pytesseract
import re

pytesseract.pytesseract.tesseract_cmd = r'C:/Users/hassa/AppData/Local/Programs/Tesseract-OCR/tesseract.exe'
os.environ["GOOGLE_API_KEY"] = "AIzaSyAWL1loBZmO2LgYkGNvpcAI912gIkL1PUo"

llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0)

def highlight_extraction(text: str) -> str:
    prompt = PromptTemplate(
        input_variables=["text"],
        template="""
        You are a helpful AI assistant. Extract the most important **action points**, **requirements**, and **commitments** from the following technical agreement.

        Write them as bullet points starting with strong verbs.

        Text:
        {text}

        Response:
        -"""
    )
    message = HumanMessage(content=prompt.format(text=text))
    return llm.invoke([message]).content.strip()

def get_flexible_splitter(text: str):
    if len(text) < 3000:
        return RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=0)
    elif len(text) < 10000:
        return RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
    else:
        return RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)

def flexible_qa(text: str , question: str, qa_chain = None) -> str:
    if len(text) <= 12000 and qa_chain is None:
        prompt = f"Answer the following question based on the text:\n\nText:\n{text}\n\nQuestion:\n{question}"
        return llm.invoke(prompt).content.strip()
    return qa_chain.run(question)

def build_qa_chain(text: str):
    if len(text) <= 12000:
        return None
    splitter = get_flexible_splitter(text)
    chunks = splitter.split_text(text)
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vectorstore = FAISS.from_texts(chunks, embedding=embeddings)
    retriever = vectorstore.as_retriever()
    return RetrievalQA.from_chain_type(llm=llm, retriever=retriever)

def extract_images_from_paragraph(para):
    images = []
    for run in para.runs:
        drawing = run._element.xpath('.//a:blip')
        for blip in drawing:
            rId = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId:
                part = para.part.related_parts[rId]
                image_bytes = part.blob
                images.append(image_bytes)
    return images

def preprocess_image(image: Image.Image) -> Image.Image:
    gray = image.convert('L')
    sharp = gray.filter(ImageFilter.SHARPEN)
    return ImageEnhance.Contrast(sharp).enhance(2.0)

def extract_tables_as_text(table):
    rows = []
    header_cells = table.rows[0].cells
    headers = [cell.text.strip() for cell in header_cells]
    rows.append(headers)
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    return rows

def prepare_doc(docx_file) -> List[dict]:
    file_bytes = docx_file.read()
    doc = Document(BytesIO(file_bytes))
    output = []
    para_idx = 0
    table_idx = 0
    diagrams_for_ascii = []

    for block in doc.element.body:
        if isinstance(block, CT_P):
            para = doc.paragraphs[para_idx]
            text = para.text.strip()

            if para.style.name.startswith("Heading") and text:
                level = int(para.style.name.replace("Heading", "").strip()) if para.style.name[-1].isdigit() else 1
                output.append({"type": "heading", "level": level, "text": text})

            elif para.style.name in ["List Bullet", "List Number"] and text:
                output.append({"type": "list_item", "text": text})

            elif text:
                output.append({"type": "paragraph", "text": text})

            images = extract_images_from_paragraph(para)
            for img_data in images:
                try:
                    image = preprocess_image(Image.open(BytesIO(img_data)))
                    raw_ocr = pytesseract.image_to_string(image, config='--oem 3 --psm 6')
                    is_potential_diagram = any(
                        kw in raw_ocr.lower()
                        for kw in ["->", "<-", "otp", "verify", "subscribe", "api", "request"]
                    )
                    if is_potential_diagram:
                        cleaned_lines = [
                            line.strip()
                            for line in raw_ocr.splitlines()
                            if line.strip() and any(x in line.lower() for x in ["->", "<-", "otp", "verify", "subscribe", "api", "request", "user", "service", "tpay", "etisalat"])
                        ]
                        diagram_text = "\n".join(cleaned_lines)
                        if diagram_text:
                            diagrams_for_ascii.append(diagram_text)
                            output.append({"type": "diagram_raw", "ocr": diagram_text})
                        else:
                            output.append({"type": "image_text", "ocr": raw_ocr})
                    else:
                        output.append({"type": "image_text", "ocr": raw_ocr})
                except Exception:
                    output.append({"type": "image_text", "ocr": "[Unreadable image]"})

            para_idx += 1

        elif isinstance(block, CT_Tbl):
            table = doc.tables[table_idx]
            rows = extract_tables_as_text(table)
            output.append({"type": "table", "rows": rows})
            table_idx += 1

    if diagrams_for_ascii:
        joined = "\n\n".join(diagrams_for_ascii)
        prompt = f"""
You are a Platform Support Engineer and System Analyst.

Your task is to convert the following extracted OCR text from a sequence diagram image into two things:

1. A clean, properly aligned **ASCII sequence diagram** using lifelines and arrows.
2. A **numbered step-by-step textual description** of the message flow.

---

Participants:
- User
- Service Provider
- TPAY
- Etisalat

---

ASCII Format Example:
User               Service Provider       TPAY                 Etisalat
  |                       |                  |                      |
  |--Subscribe----------->|                  |                      |
  |                       |--Send OTP------->|                      |
  |                       |                  |--Send SMS---------> |
  |<--Enter OTP-----------|                  |                      |
  |                       |--Verify OTP----->|                      |
  |                       |                  |--Charge-----------> |
  |                       |                  |<--Success---------- |
  |                       |<--Confirmation---|                      |
  |<--Final OK------------|                  |                      |

---

Flow Description Example:
1. User sends subscription request to Service Provider.
2. Service Provider sends OTP request to TPAY.
3. TPAY forwards OTP to Etisalat for delivery.
4. User enters OTP.
5. Service Provider verifies OTP through TPAY.
6. TPAY charges the user via Etisalat.
7. Etisalat confirms the charge.
8. Confirmation is passed back to User.

---

Text to convert:
{joined}

---

Return ONLY the ASCII diagram followed by the numbered flow description.
"""

        try:
            ascii_diagram = llm.invoke(prompt).content.strip()
        except:
            ascii_diagram = "[Gemini error: ASCII conversion failed]"
        output.append({"type": "ascii_diagram", "content": ascii_diagram})

    return output
